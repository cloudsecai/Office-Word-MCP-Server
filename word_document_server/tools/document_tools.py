"""
Document creation and manipulation tools for Word Document Server.
"""

import os
import json
from typing import List, Optional
from docx import Document

from word_document_server.utils.file_utils import (
    check_file_writeable,
    ensure_docx_extension,
    create_document_copy,
)
from word_document_server.utils.document_utils import (
    get_document_properties,
    extract_document_text,
    get_document_structure,
)
from word_document_server.core.styles import ensure_heading_style, ensure_table_style


async def create_document(
    filename: str, title: Optional[str] = None, author: Optional[str] = None
) -> str:
    """Create a new Word document (.docx format) with optional metadata.

    Creates a blank Word document with basic styles (heading and table styles) already set up.
    Only creates .docx format files - will automatically add .docx extension if not provided.

    Use this tool when:
    - Starting a new Word document from scratch
    - Need a clean document with proper styling foundations
    - Want to set document metadata (title, author) at creation time

    Args:
        filename: Name of the document to create (automatically adds .docx if missing)
        title: Optional document title (appears in document properties, not as visible content)
        author: Optional document author (appears in document properties)

    Returns:
        Success message with filename, or error message if creation fails

    Limitations:
        - Only creates .docx format (Microsoft Word 2007+)
        - Will fail if file already exists and is read-only
        - Does not add any visible content - creates blank document
    """
    filename = ensure_docx_extension(filename)

    # Check if file is writeable
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        return f"Cannot create document: {error_message}"

    try:
        doc = Document()

        # Set properties if provided
        if title:
            doc.core_properties.title = title
        if author:
            doc.core_properties.author = author

        # Ensure necessary styles exist
        ensure_heading_style(doc)
        ensure_table_style(doc)

        # Save the document
        doc.save(filename)

        return f"Document {filename} created successfully"
    except Exception as e:
        return f"Failed to create document: {str(e)}"


async def get_document_info(filename: str) -> str:
    """Get comprehensive metadata and statistics about a Word document (.docx only).

    Extracts document properties, creation/modification dates, word count, paragraph count,
    table count, and other structural information. Does not read document content.

    Use this tool when:
    - Need document metadata before processing content
    - Want to understand document size and complexity
    - Checking document properties like author, title, creation date
    - Getting quick overview without reading full content

    Args:
        filename: Path to the Word document (.docx format only)

    Returns:
        JSON formatted string containing:
        - title, author, subject, keywords
        - created/modified dates and last modified by
        - word count, paragraph count, table count, page count
        - revision number
        - error message if document doesn't exist or can't be read

    Limitations:
        - Only works with .docx format (Microsoft Word 2007+)
        - Cannot read password-protected documents
        - Word count may not match Word's count exactly for complex formatting
    """
    filename = ensure_docx_extension(filename)

    if not os.path.exists(filename):
        return f"Document {filename} does not exist"

    try:
        properties = get_document_properties(filename)
        return json.dumps(properties, indent=2)
    except Exception as e:
        return f"Failed to get document info: {str(e)}"


async def get_document_text(filename: str) -> str:
    """Extract all text content from a Word document (.docx only) with structured table formatting.

    Extracts readable text from paragraphs and tables while preserving table row/column
    relationships for optimal LLM parsing of structured data like Q&A pairs and forms.

    Use this tool when:
    - Need to read the actual content of a Word document
    - Want to process document text with AI/LLM tools
    - Converting Word content to plain text while preserving data relationships
    - Extracting text from documents with tables containing paired information

    Args:
        filename: Path to the Word document (.docx format only)

    Returns:
        Structured text with clear table formatting:
        === TABLE N ===
        Row0: | Col0: Question 1 | Col1: Answer 1 | Col2: Notes 1 |
        Row1: | Col0: Question 2 | Col1: Answer 2 | Col2: Notes 2 |
        === END TABLE ===

    Special handling:
        - Empty cells show as "(empty)"
        - Merged cells show as "(merged with above)"
        - Tables and paragraphs appear in document order
        - Clear boundaries between different tables

    Limitations:
        - Only works with .docx format (Microsoft Word 2007+)
        - Does not preserve formatting (bold, italic, fonts, colors)
        - Cannot extract images, shapes, or embedded objects
        - Cannot read password-protected documents
        - Headers/footers are not included
    """
    filename = ensure_docx_extension(filename)

    return extract_document_text(filename)


async def get_document_outline(filename: str) -> str:
    """Get structural overview of a Word document (.docx only) without full content.

    Provides a summary of document structure including paragraph previews and table
    dimensions. Useful for understanding document layout before processing full content.
    Shows first 100 characters of each paragraph and table dimensions.

    Use this tool when:
    - Need to understand document structure before full processing
    - Want to see document organization without reading everything
    - Checking if document has tables and how they're organized
    - Getting overview of content types and organization

    Args:
        filename: Path to the Word document (.docx format only)

    Returns:
        JSON formatted string containing:
        - paragraphs: array with index, preview text (first 100 chars), style name
        - tables: array with index, row count, column count, preview of first 3x3 cells

    Limitations:
        - Only works with .docx format (Microsoft Word 2007+)
        - Shows only preview text, not full content
        - Cannot read password-protected documents
        - Table previews limited to first 3 rows and 3 columns
        - Paragraph previews truncated at 100 characters
    """
    filename = ensure_docx_extension(filename)

    structure = get_document_structure(filename)
    return json.dumps(structure, indent=2)


async def list_available_documents(directory: str = ".") -> str:
    """List all Word documents (.docx files) in a specified directory.

    Scans the specified directory for .docx files and provides basic information
    about each file including size. Does not read file contents.

    Use this tool when:
    - Need to see what Word documents are available for processing
    - Want to check file sizes before processing large documents
    - Exploring a directory to find documents to work with
    - Getting inventory of available Word documents

    Args:
        directory: Directory path to scan (defaults to current directory)

    Returns:
        List of .docx files with their sizes in KB, or message if none found

    Limitations:
        - Only finds .docx format files (Microsoft Word 2007+)
        - Does not scan subdirectories
        - Cannot determine if files are readable or password-protected
        - File sizes shown in KB only
    """
    try:
        if not os.path.exists(directory):
            return f"Directory {directory} does not exist"

        docx_files = [f for f in os.listdir(directory) if f.endswith(".docx")]

        if not docx_files:
            return f"No Word documents found in {directory}"

        result = f"Found {len(docx_files)} Word documents in {directory}:\n"
        for file in docx_files:
            file_path = os.path.join(directory, file)
            size = os.path.getsize(file_path) / 1024  # KB
            result += f"- {file} ({size:.2f} KB)\n"

        return result
    except Exception as e:
        return f"Failed to list documents: {str(e)}"


async def copy_document(
    source_filename: str, destination_filename: Optional[str] = None
) -> str:
    """Create an exact copy of a Word document (.docx only) with optional rename.

    Creates a complete duplicate of a Word document including all content, formatting,
    styles, and metadata. Useful for creating backups or template variations.

    Use this tool when:
    - Creating backups before making changes to documents
    - Duplicating templates for new documents
    - Creating variations of existing documents
    - Need exact copy including all formatting and metadata

    Args:
        source_filename: Path to the source document (.docx format only)
        destination_filename: Optional new name (auto-generates if not provided)

    Returns:
        Success message with new filename, or error if copy fails

    Limitations:
        - Only works with .docx format (Microsoft Word 2007+)
        - Cannot copy password-protected documents
        - Will fail if destination already exists and is read-only
        - Source file must be readable
    """
    source_filename = ensure_docx_extension(source_filename)

    if destination_filename:
        destination_filename = ensure_docx_extension(destination_filename)

    success, message, new_path = create_document_copy(
        source_filename, destination_filename
    )
    if success:
        return message
    else:
        return f"Failed to copy document: {message}"


async def merge_documents(
    target_filename: str, source_filenames: List[str], add_page_breaks: bool = True
) -> str:
    """Merge multiple Word documents (.docx only) into a single document.

    Combines multiple Word documents by copying all paragraphs and tables from
    source documents into a new target document. Attempts to preserve basic
    formatting but may not maintain complex styles perfectly.

    Use this tool when:
    - Combining multiple related documents into one
    - Creating reports from multiple document sections
    - Consolidating document parts into final version
    - Merging template sections

    Args:
        target_filename: Name for the merged document (.docx format)
        source_filenames: List of source document paths to merge
        add_page_breaks: If True, adds page breaks between merged documents

    Returns:
        Success message with count of merged documents, or error message

    Process:
        - Creates new document with content from all sources in order
        - Copies paragraphs maintaining text and basic formatting
        - Copies tables with structure and content
        - Attempts to match styles when possible
        - Adds page breaks between documents if requested

    Limitations:
        - Only works with .docx format (Microsoft Word 2007+)
        - Complex formatting may not be preserved perfectly
        - Document order is the order in source_filenames list
        - Cannot merge password-protected documents
        - Some advanced features (headers/footers, sections) not copied
        - Style conflicts resolved by using target document styles
    """
    from word_document_server.core.tables import copy_table

    target_filename = ensure_docx_extension(target_filename)

    # Check if target file is writeable
    is_writeable, error_message = check_file_writeable(target_filename)
    if not is_writeable:
        return f"Cannot create target document: {error_message}"

    # Validate all source documents exist
    missing_files = []
    for filename in source_filenames:
        doc_filename = ensure_docx_extension(filename)
        if not os.path.exists(doc_filename):
            missing_files.append(doc_filename)

    if missing_files:
        return f"Cannot merge documents. The following source files do not exist: {', '.join(missing_files)}"

    try:
        # Create a new document for the merged result
        target_doc = Document()

        # Process each source document
        for i, filename in enumerate(source_filenames):
            doc_filename = ensure_docx_extension(filename)
            source_doc = Document(doc_filename)

            # Add page break between documents (except before the first one)
            if add_page_breaks and i > 0:
                target_doc.add_page_break()

            # Copy all paragraphs
            for paragraph in source_doc.paragraphs:
                # Create a new paragraph with the same text and style
                new_paragraph = target_doc.add_paragraph(paragraph.text)
                new_paragraph.style = target_doc.styles["Normal"]  # Default style

                # Try to match the style if possible
                try:
                    if paragraph.style and paragraph.style.name in target_doc.styles:
                        new_paragraph.style = target_doc.styles[paragraph.style.name]
                except Exception:
                    pass

                # Copy run formatting
                for i, run in enumerate(paragraph.runs):
                    if i < len(new_paragraph.runs):
                        new_run = new_paragraph.runs[i]
                        # Copy basic formatting
                        new_run.bold = run.bold
                        new_run.italic = run.italic
                        new_run.underline = run.underline
                        # Font size if specified
                        if run.font.size:
                            new_run.font.size = run.font.size

            # Copy all tables
            for table in source_doc.tables:
                copy_table(table, target_doc)

        # Save the merged document
        target_doc.save(target_filename)
        return f"Successfully merged {len(source_filenames)} documents into {target_filename}"
    except Exception as e:
        return f"Failed to merge documents: {str(e)}"
