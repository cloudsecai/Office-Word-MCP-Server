"""
Utilities for document structure and content query for Word Document Server.
"""
import os
from typing import Dict, Any
from docx import Document # Needed for Document object access

# Note: is_element_empty_util and find_text were moved here from extended_document_utils.py
# get_document_structure_details was also moved here.

def get_document_structure_details(doc_path: str) -> Dict[str, Any]:
    """
    Get detailed structure information about a Word document including 
    paragraphs, tables, styles, and run-level formatting.
    
    Args:
        doc_path: Path to the Word document
    
    Returns:
        Dictionary with comprehensive document structure details
    """
    if not os.path.exists(doc_path):
        return {"error": f"Document {doc_path} does not exist"}
    
    try:
        doc = Document(doc_path)
        structure = {
            "document_info": {
                "paragraph_count": len(doc.paragraphs),
                "table_count": len(doc.tables),
                "section_count": len(doc.sections)
            },
            "styles": [],
            "paragraphs": [],
            "tables": []
        }
        
        # Get available styles
        for style in doc.styles:
            try:
                structure["styles"].append({
                    "name": style.name,
                    "type": str(style.type),
                    "builtin": style.builtin
                })
            except (AttributeError, TypeError):
                # Skip styles that can't be accessed
                pass
        
        # Get detailed paragraph information
        for i, para in enumerate(doc.paragraphs):
            para_info = {
                "index": i,
                "text": para.text,
                "style": para.style.name if para.style else "Normal",
                "runs": []
            }
            
            # Get run-level formatting details
            for j, run in enumerate(para.runs):
                run_info = {
                    "index": j,
                    "text": run.text,
                    "bold": run.bold,
                    "italic": run.italic,
                    "underline": run.underline,
                    "font_name": run.font.name,
                    "font_size": str(run.font.size) if run.font.size else None
                }
                para_info["runs"].append(run_info)
            
            structure["paragraphs"].append(para_info)
        
        # Get detailed table information
        for i, table in enumerate(doc.tables):
            table_info = {
                "index": i,
                "rows": len(table.rows),
                "columns": len(table.columns),
                "cells": []
            }
            
            # Get cell information including merged cell detection
            for row_idx, row in enumerate(table.rows):
                row_cells = []
                for col_idx, cell in enumerate(row.cells):
                    cell_info = {
                        "row": row_idx,
                        "column": col_idx,
                        "text": cell.text,
                        "paragraphs": []
                    }
                    
                    # Get cell paragraph details
                    for para in cell.paragraphs:
                        para_details = { # Renamed from para_info to avoid conflict
                            "text": para.text,
                            "style": para.style.name if para.style else "Normal",
                            "runs": []
                        }
                        
                        # Get run details within cell paragraphs
                        for run in para.runs:
                            run_info = {
                                "text": run.text,
                                "bold": run.bold,
                                "italic": run.italic,
                                "underline": run.underline,
                                "font_name": run.font.name,
                                "font_size": str(run.font.size) if run.font.size else None
                            }
                            para_details["runs"].append(run_info)
                        
                        cell_info["paragraphs"].append(para_details)
                    
                    # Try to detect merged cells (basic implementation)
                    try:
                        # Check for horizontal merge (gridSpan)
                        tc_element = cell._tc
                        grid_span = tc_element.tcPr.gridSpan
                        if grid_span is not None:
                            cell_info["grid_span"] = grid_span.val
                        else:
                            cell_info["grid_span"] = 1
                        
                        # Check for vertical merge
                        v_merge = tc_element.tcPr.vMerge
                        if v_merge is not None:
                            cell_info["v_merge"] = v_merge.val if v_merge.val else "continue"
                        else:
                            cell_info["v_merge"] = None
                    except (AttributeError, TypeError):
                        # If we can't access merge info, set defaults
                        cell_info["grid_span"] = 1
                        cell_info["v_merge"] = None
                    
                    row_cells.append(cell_info)
                
                table_info["cells"].append(row_cells)
            
            structure["tables"].append(table_info)
        
        return structure
    except Exception as e:
        return {"error": f"Failed to get document structure details: {str(e)}"}


def find_text(doc_path: str, text_to_find: str, match_case: bool = True, whole_word: bool = False) -> Dict[str, Any]:
    """
    Find all occurrences of specific text in a Word document.
    
    Args:
        doc_path: Path to the Word document
        text_to_find: Text to search for
        match_case: Whether to perform case-sensitive search
        whole_word: Whether to match whole words only
    
    Returns:
        Dictionary with search results
    """
    if not os.path.exists(doc_path):
        return {"error": f"Document {doc_path} does not exist"}
    
    if not text_to_find:
        return {"error": "Search text cannot be empty"}
    
    try:
        doc = Document(doc_path)
        results = {
            "query": text_to_find,
            "match_case": match_case,
            "whole_word": whole_word,
            "occurrences": [],
            "total_count": 0
        }
        
        # Search in paragraphs
        for i, para in enumerate(doc.paragraphs):
            para_text_compare = para.text
            search_text_compare = text_to_find
            
            if not match_case:
                para_text_compare = para_text_compare.lower()
                search_text_compare = search_text_compare.lower()
            
            start_pos = 0
            while True:
                if whole_word:
                    # This whole_word logic could be improved with regex for more robust word boundary detection
                    # For now, sticking to split-based approach for simplicity.
                    # Re-evaluate if this becomes a frequent pain point.
                    # We need to find the index within the original casing paragraph text for context.
                    
                    # Perform search on appropriately cased text
                    # but store indices/context from original text.
                    
                    # Simple word tokenization
                    original_text_offset = 0
                    
                    temp_words_list = para.text.split() # List of words from original para.text
                    
                    # Iterate through words of the paragraph (potentially lowercased for search)
                    for word_idx, word_from_para_text_compare in enumerate(para_text_compare.split()):
                        actual_word_in_document = temp_words_list[word_idx] if word_idx < len(temp_words_list) else ""

                        if word_from_para_text_compare == search_text_compare:
                             # Try to find the start position of this word in the original paragraph text
                            try:
                                # find a word in a string, starting at a given offset
                                word_start_in_original = para.text.index(actual_word_in_document, original_text_offset)

                                results["occurrences"].append({
                                    "location_type": "paragraph",
                                    "paragraph_index": i,
                                    "position_in_paragraph": word_start_in_original, # character offset
                                    "context": para.text[:100] + ("..." if len(para.text) > 100 else "")
                                })
                                results["total_count"] += 1
                                original_text_offset = word_start_in_original + len(actual_word_in_document)
                            except ValueError:
                                # If word not found (e.g. due to split differences or prior modification), skip
                                # This could happen if split() behaves differently on cased vs uncased text with punctuation
                                pass # Or log a warning
                                
                    break # stop while loop for whole word
                else: # Substring search
                    pos = para_text_compare.find(search_text_compare, start_pos)
                    if pos == -1:
                        break
                    
                    results["occurrences"].append({
                        "location_type": "paragraph",
                        "paragraph_index": i,
                        "position_in_paragraph": pos, # character offset
                        "context": para.text[:100] + ("..." if len(para.text) > 100 else "")
                    })
                    results["total_count"] += 1
                    start_pos = pos + len(search_text_compare) # search_text_compare for length consistency
        
        # Search in tables
        for table_idx, table in enumerate(doc.tables):
            for row_idx, row in enumerate(table.rows):
                for col_idx, cell in enumerate(row.cells):
                    for para_idx, para_in_cell in enumerate(cell.paragraphs):
                        cell_para_text_compare = para_in_cell.text
                        search_text_compare = text_to_find

                        if not match_case:
                            cell_para_text_compare = cell_para_text_compare.lower()
                            search_text_compare = search_text_compare.lower()
                        
                        start_pos = 0
                        while True:
                            if whole_word:
                                temp_words_list_cell = para_in_cell.text.split()
                                original_text_offset_cell = 0
                                for word_idx, word_from_cell_para_compare in enumerate(cell_para_text_compare.split()):
                                    actual_word_in_cell_document = temp_words_list_cell[word_idx] if word_idx < len(temp_words_list_cell) else ""
                                    if word_from_cell_para_compare == search_text_compare:
                                        try:
                                            word_start_in_cell_original = para_in_cell.text.index(actual_word_in_cell_document, original_text_offset_cell)
                                            results["occurrences"].append({
                                                "location_type": "table_cell",
                                                "table_index": table_idx,
                                                "row_index": row_idx,
                                                "column_index": col_idx,
                                                "paragraph_in_cell_index": para_idx,
                                                "position_in_paragraph": word_start_in_cell_original,
                                                "context": para_in_cell.text[:100] + ("..." if len(para_in_cell.text) > 100 else "")
                                            })
                                            results["total_count"] += 1
                                            original_text_offset_cell = word_start_in_cell_original + len(actual_word_in_cell_document)
                                        except ValueError:
                                            pass # Or log
                                break # stop while loop for whole word in cell para
                            else: # Substring search
                                pos = cell_para_text_compare.find(search_text_compare, start_pos)
                                if pos == -1:
                                    break
                                
                                results["occurrences"].append({
                                    "location_type": "table_cell",
                                    "table_index": table_idx,
                                    "row_index": row_idx,
                                    "column_index": col_idx,
                                    "paragraph_in_cell_index": para_idx,
                                    "position_in_paragraph": pos,
                                    "context": para_in_cell.text[:100] + ("..." if len(para_in_cell.text) > 100 else "")
                                })
                                results["total_count"] += 1
                                start_pos = pos + len(search_text_compare)
        return results
    except Exception as e:
        return {"error": f"Failed to search for text: {str(e)}"}


def is_element_empty_util(doc_path: str, element_type: str, element_identifier: Dict[str, int]) -> Dict[str, Any]:
    """
    Check if a specific element (paragraph or table cell) is empty in a Word document.
    
    Args:
        doc_path: Path to the Word document
        element_type: Type of element ("paragraph" or "table_cell")
        element_identifier: Dictionary identifying the element
    
    Returns:
        Dictionary with operation result and boolean indicating if element is empty
    """
    if not os.path.exists(doc_path):
        return {"error": f"Document {doc_path} does not exist"}
    
    if element_type not in ["paragraph", "table_cell"]:
        return {"error": f"Invalid element_type: {element_type}. Must be 'paragraph' or 'table_cell'"}
    
    try:
        doc = Document(doc_path)
        
        if element_type == "paragraph":
            if "paragraph_index" not in element_identifier:
                return {"error": "Missing 'paragraph_index' in element_identifier for paragraph element"}
            
            paragraph_index = element_identifier["paragraph_index"]
            
            if not (0 <= paragraph_index < len(doc.paragraphs)):
                return {"error": f"Invalid paragraph index: {paragraph_index}. Document has {len(doc.paragraphs)} paragraphs."}
            
            paragraph = doc.paragraphs[paragraph_index]
            is_empty = paragraph.text.strip() == ""
            
            return {
                "success": True,
                "is_empty": is_empty,
                "element_type": element_type,
                "element_identifier": element_identifier,
                "text_content": paragraph.text if not is_empty else ""
            }
            
        elif element_type == "table_cell":
            required_keys = ["table_index", "row_index", "col_index"]
            for key in required_keys:
                if key not in element_identifier:
                    return {"error": f"Missing '{key}' in element_identifier for table_cell element"}
            
            table_index = element_identifier["table_index"]
            row_index = element_identifier["row_index"]
            col_index = element_identifier["col_index"]
            
            if not (0 <= table_index < len(doc.tables)):
                return {"error": f"Invalid table index: {table_index}. Document has {len(doc.tables)} tables."}
            table = doc.tables[table_index]
            
            if not (0 <= row_index < len(table.rows)):
                return {"error": f"Invalid row index: {row_index}. Table has {len(table.rows)} rows."}
            
            if not (0 <= col_index < len(table.columns)):
                return {"error": f"Invalid column index: {col_index}. Table has {len(table.columns)} columns."}
            
            cell = table.cell(row_index, col_index)
            is_empty = cell.text.strip() == ""
            
            return {
                "success": True,
                "is_empty": is_empty,
                "element_type": element_type,
                "element_identifier": element_identifier,
                "text_content": cell.text if not is_empty else ""
            }
        
    except Exception as e:
        return {"error": f"Failed to check if element is empty: {str(e)}"}

"""
# End of doc_structure_utils.py
""" 