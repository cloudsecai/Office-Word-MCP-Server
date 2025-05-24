"""
DocumentAnalyzer class for analyzing Word document structure and content.
"""

import os
from typing import Dict, Any, List
from docx import Document


class RunAnalyzer:
    """Helper class for analyzing run-level formatting."""

    @staticmethod
    def analyze_run(run) -> Dict[str, Any]:
        """Extract formatting information from a single run."""
        return {
            "text": run.text,
            "bold": run.bold,
            "italic": run.italic,
            "underline": run.underline,
            "font_name": run.font.name,
            "font_size": str(run.font.size) if run.font.size else None,
        }

    @staticmethod
    def analyze_runs(runs) -> List[Dict[str, Any]]:
        """Analyze a collection of runs."""
        return [
            {**RunAnalyzer.analyze_run(run), "index": idx}
            for idx, run in enumerate(runs)
        ]


class ParagraphAnalyzer:
    """Helper class for analyzing paragraph content and formatting."""

    def __init__(self):
        self.run_analyzer = RunAnalyzer()

    def analyze_paragraph(self, paragraph, index: int = None) -> Dict[str, Any]:
        """Analyze a single paragraph including its runs."""
        para_info = {
            "text": paragraph.text,
            "style": paragraph.style.name if paragraph.style else "Normal",
            "runs": self.run_analyzer.analyze_runs(paragraph.runs),
        }

        if index is not None:
            para_info["index"] = index

        return para_info

    def analyze_paragraphs(self, paragraphs) -> List[Dict[str, Any]]:
        """Analyze a collection of paragraphs."""
        return [
            self.analyze_paragraph(para, idx) for idx, para in enumerate(paragraphs)
        ]


class TableCellAnalyzer:
    """Helper class for analyzing table cell content and formatting."""

    def __init__(self):
        self.paragraph_analyzer = ParagraphAnalyzer()

    def _extract_cell_formatting(self, cell) -> Dict[str, Any]:
        """Extract merge and span information from a cell."""
        try:
            tc_element = cell._tc
            grid_span = tc_element.tcPr.gridSpan
            grid_span_val = grid_span.val if grid_span is not None else 1

            v_merge = tc_element.tcPr.vMerge
            v_merge_val = (
                v_merge.val
                if v_merge is not None and v_merge.val
                else ("continue" if v_merge is not None else None)
            )

            return {"grid_span": grid_span_val, "v_merge": v_merge_val}
        except (AttributeError, TypeError):
            return {"grid_span": 1, "v_merge": None}

    def analyze_cell(self, cell, row_idx: int, col_idx: int) -> Dict[str, Any]:
        """Analyze a single table cell."""
        cell_info = {
            "row": row_idx,
            "column": col_idx,
            "text": cell.text,
            "paragraphs": [
                self.paragraph_analyzer.analyze_paragraph(para)
                for para in cell.paragraphs
            ],
        }

        # Add formatting information
        formatting = self._extract_cell_formatting(cell)
        cell_info.update(formatting)

        return cell_info


class TableAnalyzer:
    """Helper class for analyzing table structure and content."""

    def __init__(self):
        self.cell_analyzer = TableCellAnalyzer()

    def analyze_table(self, table, index: int = None) -> Dict[str, Any]:
        """Analyze a single table including all its cells."""
        table_info = {
            "rows": len(table.rows),
            "columns": len(table.columns),
            "cells": [],
        }

        if index is not None:
            table_info["index"] = index

        # Analyze each row and cell
        for row_idx, row in enumerate(table.rows):
            row_cells = []
            for col_idx, cell in enumerate(row.cells):
                cell_info = self.cell_analyzer.analyze_cell(cell, row_idx, col_idx)
                row_cells.append(cell_info)
            table_info["cells"].append(row_cells)

        return table_info

    def analyze_tables(self, tables) -> List[Dict[str, Any]]:
        """Analyze a collection of tables."""
        return [self.analyze_table(table, idx) for idx, table in enumerate(tables)]


class StyleAnalyzer:
    """Helper class for analyzing document styles."""

    @staticmethod
    def analyze_styles(styles) -> List[Dict[str, Any]]:
        """Extract information about available styles."""
        style_info = []
        for style in styles:
            try:
                style_info.append(
                    {
                        "name": style.name,
                        "type": str(style.type),
                        "builtin": style.builtin,
                    }
                )
            except (AttributeError, TypeError):
                # Skip styles that can't be accessed
                pass
        return style_info


class DocumentAnalyzer:
    """Analyzes Word document structure and content with reusable components."""

    def __init__(self, doc_path: str):
        """Initialize with a document path."""
        self.doc_path = doc_path
        self._doc = None

        # Initialize analyzers
        self.style_analyzer = StyleAnalyzer()
        self.paragraph_analyzer = ParagraphAnalyzer()
        self.table_analyzer = TableAnalyzer()

    def _load_document(self) -> Dict[str, Any]:
        """Load and validate the document."""
        if not os.path.exists(self.doc_path):
            return {"error": f"Document {self.doc_path} does not exist"}

        try:
            self._doc = Document(self.doc_path)
            return {"success": True}
        except Exception as e:
            return {"error": f"Failed to load document: {str(e)}"}

    def get_basic_info(self) -> Dict[str, Any]:
        """Get basic document information."""
        if not self._doc:
            load_result = self._load_document()
            if "error" in load_result:
                return load_result

        return {
            "paragraph_count": len(self._doc.paragraphs),
            "table_count": len(self._doc.tables),
            "section_count": len(self._doc.sections),
        }

    def get_styles(self) -> List[Dict[str, Any]]:
        """Get available document styles."""
        if not self._doc:
            load_result = self._load_document()
            if "error" in load_result:
                return []

        return self.style_analyzer.analyze_styles(self._doc.styles)

    def get_paragraphs_analysis(self) -> List[Dict[str, Any]]:
        """Get detailed analysis of all paragraphs."""
        if not self._doc:
            load_result = self._load_document()
            if "error" in load_result:
                return []

        return self.paragraph_analyzer.analyze_paragraphs(self._doc.paragraphs)

    def get_tables_analysis(self) -> List[Dict[str, Any]]:
        """Get detailed analysis of all tables."""
        if not self._doc:
            load_result = self._load_document()
            if "error" in load_result:
                return []

        return self.table_analyzer.analyze_tables(self._doc.tables)

    def get_complete_structure(self) -> Dict[str, Any]:
        """Get complete document structure analysis."""
        if not self._doc:
            load_result = self._load_document()
            if "error" in load_result:
                return load_result

        try:
            return {
                "document_info": self.get_basic_info(),
                "styles": self.get_styles(),
                "paragraphs": self.get_paragraphs_analysis(),
                "tables": self.get_tables_analysis(),
            }
        except Exception as e:
            return {"error": f"Failed to get document structure details: {str(e)}"}

    def find_text(
        self, text_to_find: str, match_case: bool = True, whole_word: bool = False
    ) -> Dict[str, Any]:
        """Find all occurrences of specific text in the document."""
        if not self._doc:
            load_result = self._load_document()
            if "error" in load_result:
                return load_result

        if not text_to_find:
            return {"error": "Search text cannot be empty"}

        try:
            results = {
                "query": text_to_find,
                "match_case": match_case,
                "whole_word": whole_word,
                "occurrences": [],
                "total_count": 0,
            }

            # Search in paragraphs
            for i, para in enumerate(self._doc.paragraphs):
                occurrences = self._find_text_in_paragraph(
                    para, text_to_find, i, match_case, whole_word
                )
                results["occurrences"].extend(occurrences)

            # Search in table cells
            for table_idx, table in enumerate(self._doc.tables):
                for row_idx, row in enumerate(table.rows):
                    for col_idx, cell in enumerate(row.cells):
                        for para_idx, para in enumerate(cell.paragraphs):
                            occurrences = self._find_text_in_paragraph(
                                para,
                                text_to_find,
                                para_idx,
                                match_case,
                                whole_word,
                                location_context={
                                    "table_index": table_idx,
                                    "row_index": row_idx,
                                    "col_index": col_idx,
                                },
                            )
                            results["occurrences"].extend(occurrences)

            results["total_count"] = len(results["occurrences"])
            return results

        except Exception as e:
            return {"error": f"Failed to search for text: {str(e)}"}

    def _find_text_in_paragraph(
        self,
        paragraph,
        search_text: str,
        para_index: int,
        match_case: bool,
        whole_word: bool,
        location_context: Dict[str, Any] = None,
    ) -> List[Dict[str, Any]]:
        """Find text occurrences within a single paragraph."""
        para_text = paragraph.text
        search_text_compare = search_text

        if not match_case:
            para_text = para_text.lower()
            search_text_compare = search_text_compare.lower()

        occurrences = []
        start_pos = 0

        while True:
            if whole_word:
                # Simple word boundary checking
                words = para_text.split()
                for word_idx, word in enumerate(words):
                    if word == search_text_compare:
                        occurrence = {
                            "paragraph_index": para_index,
                            "word_index": word_idx,
                            "text": paragraph.text,
                            "context": " ".join(
                                words[max(0, word_idx - 2) : word_idx + 3]
                            ),
                        }
                        if location_context:
                            occurrence.update(location_context)
                        occurrences.append(occurrence)
                break
            else:
                pos = para_text.find(search_text_compare, start_pos)
                if pos == -1:
                    break

                occurrence = {
                    "paragraph_index": para_index,
                    "position": pos,
                    "text": paragraph.text,
                    "context": paragraph.text[
                        max(0, pos - 20) : pos + len(search_text) + 20
                    ],
                }
                if location_context:
                    occurrence.update(location_context)

                occurrences.append(occurrence)
                start_pos = pos + 1

        return occurrences
