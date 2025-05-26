"""
Document utility functions for Word Document Server.
"""

from typing import Dict, Any
from docx import Document
import zipfile
import xml.etree.ElementTree as ET
import logging
from collections import defaultdict


def get_document_properties(doc_path: str) -> Dict[str, Any]:
    """Get properties of a Word document."""
    import os

    if not os.path.exists(doc_path):
        return {"error": f"Document {doc_path} does not exist"}

    try:
        doc = Document(doc_path)
        core_props = doc.core_properties

        return {
            "title": core_props.title or "",
            "author": core_props.author or "",
            "subject": core_props.subject or "",
            "keywords": core_props.keywords or "",
            "created": str(core_props.created) if core_props.created else "",
            "modified": str(core_props.modified) if core_props.modified else "",
            "last_modified_by": core_props.last_modified_by or "",
            "revision": core_props.revision or 0,
            "page_count": len(doc.sections),
            "word_count": sum(
                len(paragraph.text.split()) for paragraph in doc.paragraphs
            ),
            "paragraph_count": len(doc.paragraphs),
            "table_count": len(doc.tables),
        }
    except Exception as e:
        return {"error": f"Failed to get document properties: {str(e)}"}


def get_document_text(doc_path: str) -> str:
    """Extract all text from a Word document with structured table formatting for LLM parsing, including comments and suggestions as inline tags."""
    return extract_document_text_with_comments_and_suggestions(doc_path)


def extract_document_text(doc_path: str) -> str:
    """Extract all text from a Word document with structured table formatting for LLM parsing.

    Extracts text while preserving table structure with clear row/cell boundaries
    to maintain relationships between cells (essential for Q&A pairs, forms, etc.).

    Args:
        doc_path: Path to the Word document

    Returns:
        Structured text with clear table formatting for optimal LLM parsing
    """
    import os

    if not os.path.exists(doc_path):
        return f"Document {doc_path} does not exist"

    try:
        from word_document_server.utils.document_analyzer import DocumentAnalyzer

        doc = Document(doc_path)
        analyzer = DocumentAnalyzer(doc_path)
        structure = analyzer.get_complete_structure()

        if "error" in structure:
            return structure["error"]

        text_parts = []
        table_index = 0

        # Get all body elements and process them in order
        for element in doc.element.body:
            tag_name = element.tag.split("}")[-1] if "}" in element.tag else element.tag

            if tag_name == "p":  # paragraph
                # Find the corresponding paragraph in the parsed document
                for para in doc.paragraphs:
                    if para._element == element:
                        para_text = para.text.strip()
                        if para_text:  # Only add non-empty paragraphs
                            text_parts.append(para_text)
                        break

            elif tag_name == "tbl":  # table
                if table_index < len(structure["tables"]):
                    table_text = _format_table_for_llm(
                        structure["tables"][table_index], table_index
                    )
                    text_parts.append(table_text)
                    table_index += 1

        return "\n\n".join(text_parts)

    except Exception:
        # Fallback to simple extraction if structured fails
        doc = Document(doc_path)
        text = []

        for paragraph in doc.paragraphs:
            text.append(paragraph.text)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        text.append(paragraph.text)

        return "\n".join(text)


def _format_table_for_llm(table_data: Dict[str, Any], table_index: int) -> str:
    """Format table data with clear row/cell boundaries for LLM parsing."""
    if not table_data.get("cells"):
        return f"=== TABLE {table_index + 1} (empty) ==="

    formatted_lines = [f"=== TABLE {table_index + 1} ==="]

    for row_idx, row_cells in enumerate(table_data["cells"]):
        row_parts = []
        for col_idx, cell in enumerate(row_cells):
            cell_text = cell.get("text", "").strip()
            # Handle merged cells
            if cell.get("v_merge") == "continue":
                cell_text = "(merged with above)"
            elif not cell_text:
                cell_text = "(empty)"

            row_parts.append(f"Col{col_idx}: {cell_text}")

        row_line = f"Row{row_idx}: | " + " | ".join(row_parts) + " |"
        formatted_lines.append(row_line)

    formatted_lines.append("=== END TABLE ===")
    return "\n".join(formatted_lines)


def get_document_structure(doc_path: str) -> Dict[str, Any]:
    """Get the structure of a Word document."""
    import os

    if not os.path.exists(doc_path):
        return {"error": f"Document {doc_path} does not exist"}

    try:
        doc = Document(doc_path)
        structure = {"paragraphs": [], "tables": []}

        # Get paragraphs
        for i, para in enumerate(doc.paragraphs):
            structure["paragraphs"].append(
                {
                    "index": i,
                    "text": para.text[:100] + ("..." if len(para.text) > 100 else ""),
                    "style": para.style.name if para.style else "Normal",
                }
            )

        # Get tables
        for i, table in enumerate(doc.tables):
            table_data = {
                "index": i,
                "rows": len(table.rows),
                "columns": len(table.columns),
                "preview": [],
            }

            # Get sample of table data
            max_rows = min(3, len(table.rows))
            for row_idx in range(max_rows):
                row_data = []
                max_cols = min(3, len(table.columns))
                for col_idx in range(max_cols):
                    try:
                        cell_text = table.cell(row_idx, col_idx).text
                        row_data.append(
                            cell_text[:20] + ("..." if len(cell_text) > 20 else "")
                        )
                    except IndexError:
                        row_data.append("N/A")
                table_data["preview"].append(row_data)

            structure["tables"].append(table_data)

        return structure
    except Exception as e:
        return {"error": f"Failed to get document structure: {str(e)}"}


def find_paragraph_by_text(doc, text, partial_match=False):
    """
    Find paragraphs containing specific text.

    Args:
        doc: Document object
        text: Text to search for
        partial_match: If True, matches paragraphs containing the text; if False, matches exact text

    Returns:
        List of paragraph indices that match the criteria
    """
    matching_paragraphs = []

    for i, para in enumerate(doc.paragraphs):
        if partial_match and text in para.text:
            matching_paragraphs.append(i)
        elif not partial_match and para.text == text:
            matching_paragraphs.append(i)

    return matching_paragraphs


def find_and_replace_text(doc, old_text, new_text):
    """
    Find and replace text throughout the document.

    Args:
        doc: Document object
        old_text: Text to find
        new_text: Text to replace with

    Returns:
        Number of replacements made
    """
    count = 0

    # Search in paragraphs
    for para in doc.paragraphs:
        if old_text in para.text:
            for run in para.runs:
                if old_text in run.text:
                    run.text = run.text.replace(old_text, new_text)
                    count += 1

    # Search in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if old_text in para.text:
                        for run in para.runs:
                            if old_text in run.text:
                                run.text = run.text.replace(old_text, new_text)
                                count += 1

    return count


def extract_document_text_with_comments_and_suggestions(docx_path: str) -> str:
    """
    Extracts text from a Word document (.docx), inlining comments and tracked changes (suggestions)
    using special tags for LLM consumption. Tags are defined at the top of the output if present.
    Returns a string with tags: [COMMENT], [SUGGESTION], [SUGGESTED_ADDITION], [SUGGESTED_DELETION].
    """
    TAG_DEFS = """
=== TAG DEFINITIONS ===
[COMMENT]: A user comment on the enclosed text. Format: [COMMENT ...]anchor text | comment text[/COMMENT]
[SUGGESTION]: A suggested replacement. Format: [SUGGESTION ... original=\"old text\"]new text[/SUGGESTION]
[SUGGESTED_ADDITION]: A suggested addition. Format: [SUGGESTED_ADDITION ...]added text[/SUGGESTED_ADDITION]
[SUGGESTED_DELETION]: A suggested deletion. Format: [SUGGESTED_DELETION ...]text to remove[/SUGGESTED_DELETION]
=== END TAG DEFINITIONS ===

"""
    try:
        with zipfile.ZipFile(docx_path) as docx_zip:
            # Parse main document XML
            with docx_zip.open('word/document.xml') as doc_xml:
                doc_tree = ET.parse(doc_xml)
                doc_root = doc_tree.getroot()
            # Parse comments XML (if present)
            comments = {}
            try:
                with docx_zip.open('word/comments.xml') as comments_xml:
                    comments_tree = ET.parse(comments_xml)
                    for c in comments_tree.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}comment'):
                        cid = c.attrib.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}id')
                        author = c.attrib.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}author', '')
                        date = c.attrib.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}date', '')
                        text = ''.join(t.text or '' for t in c.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t'))
                        comments[cid] = {'author': author, 'date': date, 'text': text}
            except KeyError:
                logging.info("No comments.xml found in docx.")
            except Exception as e:
                logging.error(f"Error parsing comments.xml: {e}")

            # Build a map of comment anchors (commentRangeStart/End)
            comment_starts = {}
            comment_ends = {}
            for elem in doc_root.iter():
                if elem.tag.endswith('commentRangeStart'):
                    cid = elem.attrib.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}id')
                    comment_starts[cid] = elem
                if elem.tag.endswith('commentRangeEnd'):
                    cid = elem.attrib.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}id')
                    comment_ends[cid] = elem

            # Helper to get all text in a run or element
            def get_text(e):
                return ''.join(t.text or '' for t in e.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t'))

            # Helper to get author/date from w:ins/w:del
            def get_change_metadata(e):
                author = e.attrib.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}author', '')
                date = e.attrib.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}date', '')
                return author, date

            # Walk paragraphs and tables, reconstructing text with tags
            ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
            output = []
            tags_used = set()
            para_idx = 0
            for body_child in doc_root.find('w:body', ns):
                if body_child.tag == '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p':
                    # Handle paragraph
                    para_text = ''
                    i = 0
                    runs = list(body_child)
                    while i < len(runs):
                        run = runs[i]
                        # Handle comment start
                        if run.tag.endswith('commentRangeStart'):
                            cid = run.attrib.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}id')
                            anchor_text = ''
                            j = i + 1
                            # Collect anchor text until commentRangeEnd
                            while j < len(runs):
                                if runs[j].tag.endswith('commentRangeEnd') and runs[j].attrib.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}id') == cid:
                                    break
                                anchor_text += get_text(runs[j])
                                j += 1
                            comment = comments.get(cid, {})
                            author = comment.get('author', '')
                            date = comment.get('date', '')
                            ctext = comment.get('text', '')
                            tag = f'[COMMENT id={cid} author="{author}" date="{date}"]{anchor_text} | {ctext}[/COMMENT]'
                            para_text += tag
                            tags_used.add('COMMENT')
                            i = j  # Skip to after commentRangeEnd
                        # Handle insertions (suggested additions or replacements)
                        elif run.tag.endswith('ins'):
                            author, date = get_change_metadata(run)
                            ins_text = get_text(run)
                            # Check for replacement (w:del immediately before w:ins)
                            if i > 0 and runs[i-1].tag.endswith('del'):
                                del_run = runs[i-1]
                                del_text = get_text(del_run)
                                tag = f'[SUGGESTION id={i} author="{author}" date="{date}" original="{del_text}"]{ins_text}[/SUGGESTION]'
                                tags_used.add('SUGGESTION')
                                # Remove the deletion from output (handled here)
                                para_text = para_text[:-len(del_text)] if para_text.endswith(del_text) else para_text
                            else:
                                tag = f'[SUGGESTED_ADDITION id={i} author="{author}" date="{date}"]{ins_text}[/SUGGESTED_ADDITION]'
                                tags_used.add('SUGGESTED_ADDITION')
                            para_text += tag
                        # Handle deletions (suggested deletions)
                        elif run.tag.endswith('del'):
                            author, date = get_change_metadata(run)
                            del_text = get_text(run)
                            tag = f'[SUGGESTED_DELETION id={i} author="{author}" date="{date}"]{del_text}[/SUGGESTED_DELETION]'
                            para_text += tag
                            tags_used.add('SUGGESTED_DELETION')
                        # Normal run
                        elif run.tag.endswith('r'):
                            para_text += get_text(run)
                        i += 1
                    output.append(para_text)
                    para_idx += 1
                elif body_child.tag == '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tbl':
                    # Handle table (flattened as text for now)
                    for row in body_child.findall('.//w:tr', ns):
                        row_text = []
                        for cell in row.findall('.//w:tc', ns):
                            cell_text = ''
                            for para in cell.findall('.//w:p', ns):
                                para_text = ''
                                for run in para:
                                    if run.tag.endswith('r'):
                                        para_text += get_text(run)
                                    elif run.tag.endswith('ins'):
                                        author, date = get_change_metadata(run)
                                        ins_text = get_text(run)
                                        tag = f'[SUGGESTED_ADDITION id={i} author="{author}" date="{date}"]{ins_text}[/SUGGESTED_ADDITION]'
                                        para_text += tag
                                        tags_used.add('SUGGESTED_ADDITION')
                                    elif run.tag.endswith('del'):
                                        author, date = get_change_metadata(run)
                                        del_text = get_text(run)
                                        tag = f'[SUGGESTED_DELETION id={i} author="{author}" date="{date}"]{del_text}[/SUGGESTED_DELETION]'
                                        para_text += tag
                                        tags_used.add('SUGGESTED_DELETION')
                                cell_text += para_text
                            row_text.append(cell_text)
                        output.append(' | '.join(row_text))
            result = '\n\n'.join(output)
            if tags_used:
                result = TAG_DEFS + result
            return result
    except Exception as e:
        logging.error(f"Failed to extract document text with comments/suggestions: {e}")
        return f"[ERROR] Failed to extract document text with comments/suggestions: {e}"
