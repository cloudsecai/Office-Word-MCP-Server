"""
Utilities for advanced and scoped editing operations for Word Document Server.
"""
import os
from typing import Dict, Any
from docx import Document
from docx.text.font import Font as DocxFont # For type hinting

# Function moved from extended_document_utils.py:
# search_and_replace_in_scope_util (and its helpers)

def _copy_font_properties(source_font: DocxFont, target_run):
    """Helper to copy font properties from a Font object to a Run's font."""
    if source_font is None:
        return
    target_font = target_run.font
    
    # Basic properties
    attrs_to_copy = [
        'name', 'size', 'bold', 'italic', 'underline', 'strike', 'double_strike',
        'subscript', 'superscript', 'all_caps', 'small_caps', 'shadow', 'outline',
        'emboss', 'imprint' # 'engrave' is often an alias for imprint or not directly settable.
    ]
    for attr in attrs_to_copy:
        # Check if source_font actually has the attribute because not all Font objects are fully populated
        if hasattr(source_font, attr):
            val = getattr(source_font, attr, None)
            if val is not None: # Ensure we only set actual values, not None if attr was missing
                try:
                    setattr(target_font, attr, val)
                except AttributeError: # Some properties might not be settable on newly created run fonts
                    pass # Or log a warning
    
    # Color (RGB)
    if hasattr(source_font, 'color') and source_font.color and hasattr(source_font.color, 'rgb') and source_font.color.rgb is not None:
        target_font.color.rgb = source_font.color.rgb
    
    # Highlight color
    if hasattr(source_font, 'highlight_color') and source_font.highlight_color is not None:
        try:
            target_font.highlight_color = source_font.highlight_color
        except (AttributeError, ValueError): # Handle if highlight_color is not settable or invalid enum
            pass # Or log a warning

def apply_formatted_segment(paragraph, segment_text: str, original_runs_info: list, segment_start_in_full_text: int):
    """Reconstructs a segment of text with its original run-level formatting."""
    if not segment_text:
        return

    current_run_text_buffer = ""
    current_run_font_obj = None
    
    for i, char_in_segment in enumerate(segment_text):
        original_char_absolute_pos = segment_start_in_full_text + i
        
        # Determine the font of the character at its original position
        char_original_font = paragraph.style.font # Default to paragraph style's font
        char_found_in_original_runs = False
        running_char_count_for_lookup = 0
        for run_info in original_runs_info:
            run_len = len(run_info['text'])
            if running_char_count_for_lookup <= original_char_absolute_pos < running_char_count_for_lookup + run_len:
                char_original_font = run_info['font']
                char_found_in_original_runs = True
                break
            running_char_count_for_lookup += run_len
        
        if not char_found_in_original_runs and original_runs_info: # Fallback if char is somehow out of original runs' range
             char_original_font = original_runs_info[0]['font'] # Use first run's font as a guess

        # If current_run_text_buffer is empty, start a new run
        if not current_run_text_buffer:
            current_run_text_buffer = char_in_segment
            current_run_font_obj = char_original_font
        # If font is same as current buffered run, append char
        elif current_run_font_obj and char_original_font and \
             current_run_font_obj.name == char_original_font.name and \
             current_run_font_obj.size == char_original_font.size and \
             current_run_font_obj.bold == char_original_font.bold and \
             current_run_font_obj.italic == char_original_font.italic and \
             current_run_font_obj.underline == char_original_font.underline and \
             (current_run_font_obj.color.rgb if hasattr(current_run_font_obj, 'color') and current_run_font_obj.color else None) == \
             (char_original_font.color.rgb if hasattr(char_original_font, 'color') and char_original_font.color else None) and \
             current_run_font_obj.highlight_color == char_original_font.highlight_color: # Add more properties for robust comparison
            current_run_text_buffer += char_in_segment
        # Font changed, so write out the buffered run and start a new one
        else:
            if current_run_text_buffer:
                new_run = paragraph.add_run(current_run_text_buffer)
                _copy_font_properties(current_run_font_obj, new_run)
            
            current_run_text_buffer = char_in_segment
            current_run_font_obj = char_original_font
    
    # Add the last accumulated run
    if current_run_text_buffer:
        new_run = paragraph.add_run(current_run_text_buffer)
        _copy_font_properties(current_run_font_obj, new_run)

def replace_in_paragraph_preserving_formatting(paragraph, find_str: str, replace_str: str) -> int:
    """Replaces text in a paragraph while attempting to preserve run-level formatting.
    Returns the number of replacements made in this paragraph.
    """
    replacements_made_in_para = 0
    
    original_runs_info = []
    for run in paragraph.runs:
        original_runs_info.append({
            'text': run.text,
            'font': run.font # Store the actual Font object
        })

    full_text_original = "".join(r['text'] for r in original_runs_info)
    
    if not find_str or find_str not in full_text_original:
        return 0 # No replacements

    # Clear existing runs from the paragraph (OXML level)
    paragraph._element.clear_content() 

    current_pos_in_original_text = 0
    
    while current_pos_in_original_text < len(full_text_original):
        match_start_index = full_text_original.find(find_str, current_pos_in_original_text)

        if match_start_index == -1: # No more matches
            # Add remaining text from original with its formatting
            remaining_segment = full_text_original[current_pos_in_original_text:]
            apply_formatted_segment(paragraph, remaining_segment, original_runs_info, current_pos_in_original_text)
            break # Exit while loop
        
        # Add text before the match (from current_pos up to match_start_index)
        text_before_match_segment = full_text_original[current_pos_in_original_text:match_start_index]
        apply_formatted_segment(paragraph, text_before_match_segment, original_runs_info, current_pos_in_original_text)
        
        # Add the replacement text
        # Style it like the beginning of the found text segment
        font_at_match_start = paragraph.style.font # Default
        if original_runs_info: # Ensure there are runs to get style from
            char_count_for_style_lookup = 0
            for run_info in original_runs_info:
                run_len = len(run_info['text'])
                if char_count_for_style_lookup <= match_start_index < char_count_for_style_lookup + run_len:
                    font_at_match_start = run_info['font']
                    break
                char_count_for_style_lookup += run_len
            if char_count_for_style_lookup <= match_start_index and not font_at_match_start: # If match_start is at the very end
                 font_at_match_start = original_runs_info[-1]['font']

        added_replacement_run = paragraph.add_run(replace_str)
        _copy_font_properties(font_at_match_start, added_replacement_run)
        replacements_made_in_para += 1
        
        # Move current_pos past the replaced segment
        current_pos_in_original_text = match_start_index + len(find_str)

    return replacements_made_in_para

def search_and_replace_in_scope_util(doc_path: str, find_text_str: str, replace_text_str: str, \
                                   scope_type: str, scope_identifier: Dict[str, int]) -> Dict[str, Any]:
    """
    Search and replace text within a specific scope (paragraph or table cell) 
    in a Word document, preserving run-level formatting.
    (Code adopted from previous extended_document_utils.py and significantly reworked)
    """
    total_replacements = 0

    if not os.path.exists(doc_path):
        return {"error": f"Document {doc_path} does not exist"}
    
    if not find_text_str:
        return {"error": "Find text cannot be empty"}
    
    if scope_type not in ["paragraph", "table_cell"]:
        return {"error": f"Invalid scope_type: {scope_type}. Must be 'paragraph' or 'table_cell'"}

    try:
        doc = Document(doc_path)
        
        if scope_type == "paragraph":
            if "paragraph_index" not in scope_identifier:
                return {"error": "Missing 'paragraph_index' in scope_identifier for paragraph scope"}
            paragraph_index = scope_identifier["paragraph_index"]
            if not (0 <= paragraph_index < len(doc.paragraphs)):
                return {"error": f"Invalid paragraph index: {paragraph_index}. Document has {len(doc.paragraphs)} paragraphs."}
            
            target_paragraph = doc.paragraphs[paragraph_index]
            total_replacements += replace_in_paragraph_preserving_formatting(target_paragraph, find_text_str, replace_text_str)
            
        elif scope_type == "table_cell":
            required_keys = ["table_index", "row_index", "col_index"]
            for key in required_keys:
                if key not in scope_identifier:
                    return {"error": f"Missing '{key}' in scope_identifier for table_cell scope"}
            
            table_index = scope_identifier["table_index"]
            row_index = scope_identifier["row_index"]
            col_index = scope_identifier["col_index"]

            if not (0 <= table_index < len(doc.tables)):
                return {"error": f"Invalid table index: {table_index}. Document has {len(doc.tables)} tables."}
            table = doc.tables[table_index]

            if not (0 <= row_index < len(table.rows)):
                return {"error": f"Invalid row index: {row_index}. Table has {len(table.rows)} rows."}
            
            if not (0 <= col_index < len(table.columns)):
                 return {"error": f"Invalid column index: {col_index}. Table has {len(table.columns)} columns."}
            
            cell = table.cell(row_index, col_index)
            for p_in_cell in cell.paragraphs:
                total_replacements += replace_in_paragraph_preserving_formatting(p_in_cell, find_text_str, replace_text_str)
        
        if total_replacements > 0:
            doc.save(doc_path)
        
        return {
            "success": True,
            "message": f"Replaced {total_replacements} occurrence(s) of '{find_text_str}' with '{replace_text_str}' in {scope_type}",
            "replacements_made": total_replacements,
            "find_text": find_text_str,
            "replace_text": replace_text_str,
            "scope_type": scope_type,
            "scope_identifier": scope_identifier
        }
        
    except Exception as e:
        # Consider more specific error logging here if possible
        return {"error": f"Failed to search and replace in scope: {str(e)}"}

"""
End of editing_utils.py
""" 