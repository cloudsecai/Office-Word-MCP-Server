"""
FormattedEditor class for text replacement operations while preserving formatting.
"""

import os
from typing import Dict, Any, List, Optional
from dataclasses import dataclass
from docx import Document
from docx.text.font import Font as DocxFont


@dataclass
class ScopeLocation:
    """Data class representing a scope for text operations."""

    scope_type: str  # "paragraph" or "table_cell"
    paragraph_index: Optional[int] = None
    table_index: Optional[int] = None
    row_index: Optional[int] = None
    col_index: Optional[int] = None

    def __str__(self) -> str:
        if self.scope_type == "paragraph":
            return f"Paragraph[{self.paragraph_index}]"
        else:
            return f"Table[{self.table_index}] Cell[{self.row_index}, {self.col_index}]"


@dataclass
class RunInfo:
    """Data class representing run formatting information."""

    text: str
    font: DocxFont


class FontFormatter:
    """Helper class for managing font properties and formatting."""

    @staticmethod
    def copy_font_properties(source_font: DocxFont, target_run) -> None:
        """Copy font properties from source to target run."""
        if source_font is None:
            return

        target_font = target_run.font

        # Basic properties to copy
        attrs_to_copy = [
            "name",
            "size",
            "bold",
            "italic",
            "underline",
            "strike",
            "double_strike",
            "subscript",
            "superscript",
            "all_caps",
            "small_caps",
            "shadow",
            "outline",
            "emboss",
            "imprint",
        ]

        for attr in attrs_to_copy:
            if hasattr(source_font, attr):
                val = getattr(source_font, attr, None)
                if val is not None:
                    try:
                        setattr(target_font, attr, val)
                    except AttributeError:
                        # Some properties might not be settable
                        pass

        # Handle color
        if (
            hasattr(source_font, "color")
            and source_font.color
            and hasattr(source_font.color, "rgb")
            and source_font.color.rgb is not None
        ):
            target_font.color.rgb = source_font.color.rgb

        # Handle highlight color
        if (
            hasattr(source_font, "highlight_color")
            and source_font.highlight_color is not None
        ):
            try:
                target_font.highlight_color = source_font.highlight_color
            except (AttributeError, ValueError):
                pass

    @staticmethod
    def fonts_are_equivalent(font1: DocxFont, font2: DocxFont) -> bool:
        """Check if two fonts have equivalent formatting."""
        if not font1 or not font2:
            return font1 == font2

        # Compare key properties
        return (
            font1.name == font2.name
            and font1.size == font2.size
            and font1.bold == font2.bold
            and font1.italic == font2.italic
            and font1.underline == font2.underline
            and FontFormatter._get_color_rgb(font1)
            == FontFormatter._get_color_rgb(font2)
            and font1.highlight_color == font2.highlight_color
        )

    @staticmethod
    def _get_color_rgb(font: DocxFont):
        """Safely get RGB color from font."""
        try:
            return font.color.rgb if hasattr(font, "color") and font.color else None
        except (AttributeError, TypeError):
            return None


class ParagraphFormatter:
    """Helper class for managing paragraph-level formatting operations."""

    def __init__(self):
        self.font_formatter = FontFormatter()

    def extract_run_info(self, paragraph) -> List[RunInfo]:
        """Extract run information from a paragraph."""
        return [RunInfo(text=run.text, font=run.font) for run in paragraph.runs]

    def apply_formatted_segment(
        self,
        paragraph,
        segment_text: str,
        original_runs: List[RunInfo],
        segment_start_pos: int,
    ) -> None:
        """Apply a text segment with preserved formatting."""
        if not segment_text:
            return

        current_run_buffer = ""
        current_font = None

        for i, char in enumerate(segment_text):
            char_pos = segment_start_pos + i
            char_font = self._get_font_for_position(char_pos, original_runs, paragraph)

            # Start new run or continue existing one
            if not current_run_buffer:
                current_run_buffer = char
                current_font = char_font
            elif self.font_formatter.fonts_are_equivalent(current_font, char_font):
                current_run_buffer += char
            else:
                # Font changed - flush current run and start new one
                self._add_formatted_run(paragraph, current_run_buffer, current_font)
                current_run_buffer = char
                current_font = char_font

        # Add final run
        if current_run_buffer:
            self._add_formatted_run(paragraph, current_run_buffer, current_font)

    def _get_font_for_position(
        self, position: int, original_runs: List[RunInfo], paragraph
    ) -> DocxFont:
        """Get the font that should be used for a character at a specific position."""
        current_pos = 0

        for run_info in original_runs:
            run_length = len(run_info.text)
            if current_pos <= position < current_pos + run_length:
                return run_info.font
            current_pos += run_length

        # Fallback to paragraph style or first run
        if original_runs:
            return original_runs[0].font
        return paragraph.style.font

    def _add_formatted_run(self, paragraph, text: str, font: DocxFont) -> None:
        """Add a new run to the paragraph with specific formatting."""
        new_run = paragraph.add_run(text)
        self.font_formatter.copy_font_properties(font, new_run)


class TextReplacer:
    """Core class for performing text replacement while preserving formatting."""

    def __init__(self):
        self.paragraph_formatter = ParagraphFormatter()

    def replace_in_paragraph(self, paragraph, find_text: str, replace_text: str) -> int:
        """Replace text in a paragraph while preserving formatting."""
        if not find_text or find_text not in paragraph.text:
            return 0

        # Extract original run information
        original_runs = self.paragraph_formatter.extract_run_info(paragraph)
        full_text = paragraph.text

        # Clear existing content
        paragraph._element.clear_content()

        replacements_made = 0
        current_pos = 0

        while current_pos < len(full_text):
            match_pos = full_text.find(find_text, current_pos)

            if match_pos == -1:
                # Add remaining text
                remaining_text = full_text[current_pos:]
                self.paragraph_formatter.apply_formatted_segment(
                    paragraph, remaining_text, original_runs, current_pos
                )
                break

            # Add text before match
            before_match = full_text[current_pos:match_pos]
            self.paragraph_formatter.apply_formatted_segment(
                paragraph, before_match, original_runs, current_pos
            )

            # Add replacement text with formatting from match position
            match_font = self.paragraph_formatter._get_font_for_position(
                match_pos, original_runs, paragraph
            )
            self.paragraph_formatter._add_formatted_run(
                paragraph, replace_text, match_font
            )

            replacements_made += 1
            current_pos = match_pos + len(find_text)

        return replacements_made


class FormattedEditor:
    """Main class for performing formatted text editing operations."""

    def __init__(self, doc_path: str):
        """Initialize with a document path."""
        self.doc_path = doc_path
        self._doc = None
        self.text_replacer = TextReplacer()

    def _load_document(self) -> Dict[str, Any]:
        """Load and validate the document."""
        if not os.path.exists(self.doc_path):
            return {"error": f"Document {self.doc_path} does not exist"}

        try:
            self._doc = Document(self.doc_path)
            return {"success": True}
        except Exception as e:
            return {"error": f"Failed to load document: {str(e)}"}

    def _validate_scope(self, scope: ScopeLocation) -> Dict[str, Any]:
        """Validate that the scope location is valid."""
        if not self._doc:
            load_result = self._load_document()
            if "error" in load_result:
                return load_result

        if scope.scope_type == "paragraph":
            if scope.paragraph_index is None:
                return {"error": "Missing paragraph_index for paragraph scope"}
            if not (0 <= scope.paragraph_index < len(self._doc.paragraphs)):
                return {
                    "error": f"Invalid paragraph index: {scope.paragraph_index}. Document has {len(self._doc.paragraphs)} paragraphs."
                }
            return {
                "success": True,
                "target": self._doc.paragraphs[scope.paragraph_index],
            }

        elif scope.scope_type == "table_cell":
            required_attrs = ["table_index", "row_index", "col_index"]
            for attr in required_attrs:
                if getattr(scope, attr, None) is None:
                    return {"error": f"Missing {attr} for table_cell scope"}

            if not (0 <= scope.table_index < len(self._doc.tables)):
                return {
                    "error": f"Invalid table index: {scope.table_index}. Document has {len(self._doc.tables)} tables."
                }

            table = self._doc.tables[scope.table_index]
            if not (0 <= scope.row_index < len(table.rows)):
                return {
                    "error": f"Invalid row index: {scope.row_index}. Table has {len(table.rows)} rows."
                }

            if not (0 <= scope.col_index < len(table.columns)):
                return {
                    "error": f"Invalid column index: {scope.col_index}. Table has {len(table.columns)} columns."
                }

            try:
                cell = table.cell(scope.row_index, scope.col_index)
                return {"success": True, "target": cell}
            except IndexError:
                return {"error": f"Failed to access cell at {scope}"}

        else:
            return {
                "error": f"Invalid scope_type: {scope.scope_type}. Must be 'paragraph' or 'table_cell'"
            }

    def search_and_replace_in_scope(
        self, find_text: str, replace_text: str, scope: ScopeLocation
    ) -> Dict[str, Any]:
        """Search and replace text within a specific scope."""
        if not find_text:
            return {"error": "Find text cannot be empty"}

        validation_result = self._validate_scope(scope)
        if "error" in validation_result:
            return validation_result

        try:
            target = validation_result["target"]
            total_replacements = 0

            if scope.scope_type == "paragraph":
                total_replacements = self.text_replacer.replace_in_paragraph(
                    target, find_text, replace_text
                )

            elif scope.scope_type == "table_cell":
                # Replace in all paragraphs within the cell
                for paragraph in target.paragraphs:
                    total_replacements += self.text_replacer.replace_in_paragraph(
                        paragraph, find_text, replace_text
                    )

            self._doc.save(self.doc_path)

            return {
                "success": True,
                "message": f"Replaced '{find_text}' with '{replace_text}' in {scope}",
                "scope": str(scope),
                "replacements_made": total_replacements,
                "find_text": find_text,
                "replace_text": replace_text,
            }

        except Exception as e:
            return {"error": f"Failed to perform search and replace: {str(e)}"}
