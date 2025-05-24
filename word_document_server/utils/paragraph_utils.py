"""
Utilities for paragraph-specific operations for Word Document Server.
"""

import os
from typing import Dict, Any, Optional
from docx import Document
from docx.oxml import OxmlElement  # For insert_paragraph_after_index_util
from docx.oxml.ns import qn  # For insert_paragraph_after_index_util

# Functions moved from extended_document_utils.py:
# get_paragraph_text, set_paragraph_text_util, insert_paragraph_after_index_util


def get_paragraph_text(doc_path: str, paragraph_index: int) -> Dict[str, Any]:
    """
    Get text from a specific paragraph in a Word document.

    Args:
        doc_path: Path to the Word document
        paragraph_index: Index of the paragraph to extract (0-based)

    Returns:
        Dictionary with paragraph text and metadata
    """
    if not os.path.exists(doc_path):
        return {"error": f"Document {doc_path} does not exist"}

    try:
        doc = Document(doc_path)

        if not (0 <= paragraph_index < len(doc.paragraphs)):
            return {
                "error": f"Invalid paragraph index: {paragraph_index}. Document has {len(doc.paragraphs)} paragraphs."
            }

        paragraph = doc.paragraphs[paragraph_index]

        return {
            "index": paragraph_index,
            "text": paragraph.text,
            "style": paragraph.style.name if paragraph.style else "Normal",
            "is_heading": paragraph.style.name.startswith("Heading")
            if paragraph.style
            else False,
        }
    except Exception as e:
        return {"error": f"Failed to get paragraph text: {str(e)}"}


def set_paragraph_text_util(
    doc_path: str,
    paragraph_index: int,
    new_text: str,
    style_to_apply: Optional[str] = None,
) -> Dict[str, Any]:
    """
    Set text in a specific paragraph in a Word document.

    Args:
        doc_path: Path to the Word document
        paragraph_index: Index of the paragraph (0-based)
        new_text: New text to set
        style_to_apply: Optional style to apply to the paragraph

    Returns:
        Dictionary with operation result
    """
    if not os.path.exists(doc_path):
        return {"error": f"Document {doc_path} does not exist"}

    try:
        doc = Document(doc_path)

        if not (0 <= paragraph_index < len(doc.paragraphs)):
            return {
                "error": f"Invalid paragraph index: {paragraph_index}. Document has {len(doc.paragraphs)} paragraphs."
            }

        paragraph = doc.paragraphs[paragraph_index]
        paragraph.text = new_text

        if style_to_apply:
            try:
                paragraph.style = style_to_apply
            except KeyError:  # Python-docx raises KeyError if style doesn't exist
                return {"error": f"Style '{style_to_apply}' not found in document"}

        doc.save(doc_path)

        return {
            "success": True,
            "message": f"Text set in paragraph {paragraph_index}",
            "paragraph_index": paragraph_index,
            "text_set": new_text,
            "style_applied": style_to_apply,
        }
    except Exception as e:
        return {"error": f"Failed to set paragraph text: {str(e)}"}


def insert_paragraph_after_index_util(
    doc_path: str,
    target_paragraph_index: int,
    text_to_insert: str,
    style_to_apply: Optional[str] = None,
) -> Dict[str, Any]:
    """
    Insert a new paragraph after a specific paragraph index in a Word document.

    Args:
        doc_path: Path to the Word document
        target_paragraph_index: Index of the paragraph after which to insert (0-based)
        text_to_insert: Text for the new paragraph
        style_to_apply: Optional style to apply to the new paragraph

    Returns:
        Dictionary with operation result
    """
    if not os.path.exists(doc_path):
        return {"error": f"Document {doc_path} does not exist"}

    try:
        doc = Document(doc_path)

        if not (0 <= target_paragraph_index < len(doc.paragraphs)):
            return {
                "error": f"Invalid paragraph index: {target_paragraph_index}. Document has {len(doc.paragraphs)} paragraphs."
            }

        target_paragraph = doc.paragraphs[target_paragraph_index]

        new_p_oxml = OxmlElement("w:p")
        new_r_oxml = OxmlElement("w:r")
        new_t_oxml = OxmlElement("w:t")
        new_t_oxml.text = text_to_insert
        new_r_oxml.append(new_t_oxml)
        new_p_oxml.append(new_r_oxml)

        if style_to_apply:
            pPr = new_p_oxml.get_or_add_pPr()
            pStyle = OxmlElement("w:pStyle")
            pStyle.set(qn("w:val"), style_to_apply)
            pPr.append(pStyle)

        target_paragraph._element.addnext(new_p_oxml)

        doc.save(doc_path)

        # To return the index of the newly inserted paragraph, it would be target_paragraph_index + 1
        # However, this can be complex if other operations happen. For now, confirming success.
        return {
            "success": True,
            "message": f"Paragraph inserted after index {target_paragraph_index}",
            "target_paragraph_index": target_paragraph_index,
            "text_inserted": text_to_insert,
            "style_applied": style_to_apply,
            "new_paragraph_approx_index": target_paragraph_index + 1,
        }
    except Exception as e:
        return {"error": f"Failed to insert paragraph: {str(e)}"}


"""
# End of paragraph_utils.py
"""
