"""
TableManager class for encapsulating table operations in Word documents.
"""

import os
from typing import Dict, Any, Optional
from dataclasses import dataclass
from docx import Document


@dataclass
class CellLocation:
    """Data class representing a table cell location."""

    table_index: int
    row_index: int
    col_index: int

    def __str__(self) -> str:
        return f"Table[{self.table_index}] Cell[{self.row_index}, {self.col_index}]"


class TableManager:
    """Manages table operations in Word documents with proper validation and error handling."""

    def __init__(self, doc_path: str):
        """Initialize with a document path."""
        self.doc_path = doc_path
        self._doc = None

    def _load_document(self) -> Dict[str, Any]:
        """Load and validate the document. Returns error dict if failed."""
        if not os.path.exists(self.doc_path):
            return {"error": f"Document {self.doc_path} does not exist"}

        try:
            self._doc = Document(self.doc_path)
            return {"success": True}
        except Exception as e:
            return {"error": f"Failed to load document: {str(e)}"}

    def _validate_cell_location(self, location: CellLocation) -> Dict[str, Any]:
        """Validate that the cell location is valid for the current document."""
        if not self._doc:
            load_result = self._load_document()
            if "error" in load_result:
                return load_result

        if not (0 <= location.table_index < len(self._doc.tables)):
            return {
                "error": f"Invalid table index: {location.table_index}. Document has {len(self._doc.tables)} tables."
            }

        table = self._doc.tables[location.table_index]

        if not (0 <= location.row_index < len(table.rows)):
            return {
                "error": f"Invalid row index: {location.row_index}. Table has {len(table.rows)} rows."
            }

        if not (0 <= location.col_index < len(table.columns)):
            return {
                "error": f"Invalid column index: {location.col_index}. Table has {len(table.columns)} columns."
            }

        return {"success": True, "table": table}

    def _analyze_cell_formatting(self, cell) -> Dict[str, Any]:
        """Extract formatting information from a cell."""
        try:
            tc_element = cell._tc
            grid_span_el = tc_element.tcPr.gridSpan
            grid_span = grid_span_el.val if grid_span_el is not None else 1

            v_merge_el = tc_element.tcPr.vMerge
            v_merge = (
                v_merge_el.val
                if v_merge_el is not None and v_merge_el.val
                else ("continue" if v_merge_el is not None else None)
            )

            return {"grid_span": grid_span, "v_merge": v_merge}
        except (AttributeError, TypeError):
            return {"grid_span": 1, "v_merge": None}

    def _analyze_paragraph_runs(self, paragraph) -> list:
        """Extract run information from a paragraph."""
        runs = []
        for run_idx, run in enumerate(paragraph.runs):
            run_info = {
                "index": run_idx,
                "text": run.text,
                "bold": run.bold,
                "italic": run.italic,
                "underline": run.underline,
                "font_name": run.font.name,
                "font_size": str(run.font.size) if run.font.size else None,
            }
            runs.append(run_info)
        return runs

    def get_cell_content(self, location: CellLocation) -> Dict[str, Any]:
        """Get detailed content from a specific table cell."""
        validation_result = self._validate_cell_location(location)
        if "error" in validation_result:
            return validation_result

        try:
            table = validation_result["table"]
            cell = table.cell(location.row_index, location.col_index)

            cell_content = {
                "location": str(location),
                "table_index": location.table_index,
                "row_index": location.row_index,
                "col_index": location.col_index,
                "text": cell.text,
                "paragraphs": [],
            }

            # Analyze paragraphs
            for para_idx, para in enumerate(cell.paragraphs):
                para_info = {
                    "index": para_idx,
                    "text": para.text,
                    "style": para.style.name if para.style else "Normal",
                    "runs": self._analyze_paragraph_runs(para),
                }
                cell_content["paragraphs"].append(para_info)

            # Add formatting information
            formatting = self._analyze_cell_formatting(cell)
            cell_content.update(formatting)

            return cell_content

        except IndexError:
            return {
                "error": f"Failed to access {location}. Index might be out of bounds for this specific row/column configuration."
            }
        except Exception as e:
            return {"error": f"Failed to get table cell content: {str(e)}"}

    def set_cell_text(
        self,
        location: CellLocation,
        text: str,
        clear_existing: bool = True,
        style: Optional[str] = None,
    ) -> Dict[str, Any]:
        """Set text in a specific table cell."""
        validation_result = self._validate_cell_location(location)
        if "error" in validation_result:
            return validation_result

        try:
            table = validation_result["table"]
            cell = table.cell(location.row_index, location.col_index)

            target_paragraph = None
            if clear_existing:
                cell.text = text
                if cell.paragraphs:
                    target_paragraph = cell.paragraphs[0]
            else:
                target_paragraph = cell.add_paragraph(text)

            if style and target_paragraph:
                try:
                    target_paragraph.style = style
                except KeyError:
                    return {"error": f"Style '{style}' not found in document"}

            self._doc.save(self.doc_path)

            return {
                "success": True,
                "message": f"Text set in {location}",
                "location": str(location),
                "text_set": text,
                "style_applied": style if target_paragraph and style else None,
            }

        except IndexError:
            return {
                "error": f"Failed to access {location}. Index might be out of bounds."
            }
        except Exception as e:
            return {"error": f"Failed to set table cell text: {str(e)}"}

    def clear_cell_content(self, location: CellLocation) -> Dict[str, Any]:
        """Clear all content from a specific table cell."""
        validation_result = self._validate_cell_location(location)
        if "error" in validation_result:
            return validation_result

        try:
            table = validation_result["table"]
            cell = table.cell(location.row_index, location.col_index)
            cell.text = ""  # Clears all paragraphs and adds a single empty one

            self._doc.save(self.doc_path)

            return {
                "success": True,
                "message": f"Content cleared from {location}",
                "location": str(location),
            }

        except IndexError:
            return {
                "error": f"Failed to access {location}. Index might be out of bounds."
            }
        except Exception as e:
            return {"error": f"Failed to clear table cell content: {str(e)}"}

    def add_paragraph_to_cell(
        self, location: CellLocation, text: str, style: Optional[str] = None
    ) -> Dict[str, Any]:
        """Add a new paragraph to a specific table cell."""
        validation_result = self._validate_cell_location(location)
        if "error" in validation_result:
            return validation_result

        try:
            table = validation_result["table"]
            cell = table.cell(location.row_index, location.col_index)
            new_paragraph = cell.add_paragraph(text)

            if style:
                try:
                    new_paragraph.style = style
                except KeyError:
                    return {"error": f"Style '{style}' not found in document"}

            self._doc.save(self.doc_path)

            return {
                "success": True,
                "message": f"Paragraph added to {location}",
                "location": str(location),
                "paragraph_text": text,
                "style_applied": style,
            }

        except IndexError:
            return {
                "error": f"Failed to access {location}. Index might be out of bounds."
            }
        except Exception as e:
            return {"error": f"Failed to add paragraph to table cell: {str(e)}"}
