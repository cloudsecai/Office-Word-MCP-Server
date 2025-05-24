"""
Utilities for table-specific operations for Word Document Server.
"""
import os
from typing import Dict, Any, Optional
from docx import Document

# Functions moved from extended_document_utils.py:
# get_table_cell_content, set_table_cell_text_util,
# clear_table_cell_content_util, add_paragraph_to_table_cell_util

def get_table_cell_content(doc_path: str, table_index: int, row_index: int, col_index: int) -> Dict[str, Any]:
    """
    Get detailed content from a specific table cell.
    
    Args:
        doc_path: Path to the Word document
        table_index: Index of the table (0-based)
        row_index: Index of the row (0-based)
        col_index: Index of the column (0-based)
    
    Returns:
        Dictionary with cell content and formatting details
    """
    if not os.path.exists(doc_path):
        return {"error": f"Document {doc_path} does not exist"}
    
    try:
        doc = Document(doc_path)
        
        if not (0 <= table_index < len(doc.tables)):
            return {"error": f"Invalid table index: {table_index}. Document has {len(doc.tables)} tables."}
        table = doc.tables[table_index]
        
        if not (0 <= row_index < len(table.rows)):
            return {"error": f"Invalid row index: {row_index}. Table has {len(table.rows)} rows."}
        
        # Ensure col_index is valid for the specific row if columns can be variable
        # For simplicity, python-docx table.cell() handles ragged tables gracefully by raising IndexError.
        if not (0 <= col_index < len(table.columns)):
             return {"error": f"Invalid column index: {col_index}. Table has {len(table.columns)} columns (may vary by row)."}

        cell = table.cell(row_index, col_index)
        
        cell_content = {
            "table_index": table_index,
            "row_index": row_index,
            "col_index": col_index,
            "text": cell.text,
            "paragraphs": []
        }
        
        for para_idx, para in enumerate(cell.paragraphs):
            para_info = {
                "index": para_idx,
                "text": para.text,
                "style": para.style.name if para.style else "Normal",
                "runs": []
            }
            for run_idx, run in enumerate(para.runs):
                run_info = {
                    "index": run_idx,
                    "text": run.text,
                    "bold": run.bold,
                    "italic": run.italic,
                    "underline": run.underline,
                    "font_name": run.font.name,
                    "font_size": str(run.font.size) if run.font.size else None
                }
                para_info["runs"].append(run_info)
            cell_content["paragraphs"].append(para_info)
        
        try:
            tc_element = cell._tc
            grid_span_el = tc_element.tcPr.gridSpan
            cell_content["grid_span"] = grid_span_el.val if grid_span_el is not None else 1
            v_merge_el = tc_element.tcPr.vMerge
            cell_content["v_merge"] = v_merge_el.val if v_merge_el is not None and v_merge_el.val else (
                "continue" if v_merge_el is not None else None
            )
        except (AttributeError, TypeError):
            cell_content["grid_span"] = 1
            cell_content["v_merge"] = None
            
        return cell_content
    except IndexError: # Raised by table.cell(row_idx, col_idx) if indices are out of bounds for ragged table
        return {"error": f"Failed to access cell ({row_index}, {col_index}) in table {table_index}. Index might be out of bounds for this specific row/column configuration."}
    except Exception as e:
        return {"error": f"Failed to get table cell content: {str(e)}"}


def set_table_cell_text_util(doc_path: str, table_index: int, row_index: int, col_index: int, 
                            text_to_set: str, clear_existing_content: bool = True, 
                            paragraph_style: Optional[str] = None) -> Dict[str, Any]:
    """
    Set text in a specific table cell in a Word document.
    (Code adopted from previous extended_document_utils.py and refined)
    """
    if not os.path.exists(doc_path):
        return {"error": f"Document {doc_path} does not exist"}
    
    try:
        doc = Document(doc_path)
        
        if not (0 <= table_index < len(doc.tables)):
            return {"error": f"Invalid table index: {table_index}. Document has {len(doc.tables)} tables."}
        table = doc.tables[table_index]
        
        if not (0 <= row_index < len(table.rows)):
            return {"error": f"Invalid row index: {row_index}. Table has {len(table.rows)} rows."}
        
        if not (0 <= col_index < len(table.columns)):
             return {"error": f"Invalid column index: {col_index}. Table has {len(table.columns)} columns."}

        cell = table.cell(row_index, col_index)
        
        target_paragraph = None
        if clear_existing_content:
            cell.text = text_to_set
            if cell.paragraphs:
                target_paragraph = cell.paragraphs[0]
        else:
            target_paragraph = cell.add_paragraph(text_to_set)
        
        if paragraph_style and target_paragraph:
            try:
                target_paragraph.style = paragraph_style
            except KeyError:
                return {"error": f"Style '{paragraph_style}' not found in document"}
        
        doc.save(doc_path)
        return {
            "success": True,
            "message": f"Text set in table {table_index}, row {row_index}, column {col_index}",
            "table_index": table_index, "row_index": row_index, "col_index": col_index,
            "text_set": text_to_set,
            "style_applied": paragraph_style if target_paragraph and paragraph_style else None
        }
    except IndexError:
        return {"error": f"Failed to access cell ({row_index}, {col_index}) in table {table_index}. Index might be out of bounds."}
    except Exception as e:
        return {"error": f"Failed to set table cell text: {str(e)}"}


def clear_table_cell_content_util(doc_path: str, table_index: int, row_index: int, col_index: int) -> Dict[str, Any]:
    """
    Clear all content from a specific table cell in a Word document.
    """
    if not os.path.exists(doc_path):
        return {"error": f"Document {doc_path} does not exist"}
    
    try:
        doc = Document(doc_path)
        if not (0 <= table_index < len(doc.tables)):
            return {"error": f"Invalid table index: {table_index}. Document has {len(doc.tables)} tables."}
        table = doc.tables[table_index]
        
        if not (0 <= row_index < len(table.rows)):
            return {"error": f"Invalid row index: {row_index}. Table has {len(table.rows)} rows."}

        if not (0 <= col_index < len(table.columns)):
             return {"error": f"Invalid column index: {col_index}. Table has {len(table.columns)} columns."}
        
        cell = table.cell(row_index, col_index)
        cell.text = '' # Clears all paragraphs and adds a single empty one.
        doc.save(doc_path)
        return {
            "success": True,
            "message": f"Content cleared from table {table_index}, row {row_index}, column {col_index}",
            "table_index": table_index, "row_index": row_index, "col_index": col_index
        }
    except IndexError:
        return {"error": f"Failed to access cell ({row_index}, {col_index}) in table {table_index}. Index might be out of bounds."}
    except Exception as e:
        return {"error": f"Failed to clear table cell content: {str(e)}"}


def add_paragraph_to_table_cell_util(doc_path: str, table_index: int, row_index: int, col_index: int, 
                                    paragraph_text: str, paragraph_style: Optional[str] = None) -> Dict[str, Any]:
    """
    Add a new paragraph to a specific table cell in a Word document.
    """
    if not os.path.exists(doc_path):
        return {"error": f"Document {doc_path} does not exist"}
    
    try:
        doc = Document(doc_path)
        if not (0 <= table_index < len(doc.tables)):
            return {"error": f"Invalid table index: {table_index}. Document has {len(doc.tables)} tables."}
        table = doc.tables[table_index]

        if not (0 <= row_index < len(table.rows)):
            return {"error": f"Invalid row index: {row_index}. Table has {len(table.rows)} rows."}
        
        if not (0 <= col_index < len(table.columns)):
             return {"error": f"Invalid column index: {col_index}. Table has {len(table.columns)} columns."}

        cell = table.cell(row_index, col_index)
        new_paragraph = cell.add_paragraph(paragraph_text)
        
        if paragraph_style:
            try:
                new_paragraph.style = paragraph_style
            except KeyError:
                return {"error": f"Style '{paragraph_style}' not found in document"}
        
        doc.save(doc_path)
        return {
            "success": True,
            "message": f"Paragraph added to table {table_index}, row {row_index}, column {col_index}",
            "table_index": table_index, "row_index": row_index, "col_index": col_index,
            "paragraph_text": paragraph_text,
            "style_applied": paragraph_style
        }
    except IndexError:
        return {"error": f"Failed to access cell ({row_index}, {col_index}) in table {table_index}. Index might be out of bounds."}
    except Exception as e:
        return {"error": f"Failed to add paragraph to table cell: {str(e)}"}

"""
# End of table_utils.py
""" 